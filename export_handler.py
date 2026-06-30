"""
Export Handler Module
Handles exporting timetables to CSV and Word formats
"""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime, timedelta
import re

class TimetableExporter:
    """Handles timetable export to various formats"""

    @staticmethod
    def _parse_time_range(time_range: str) -> Optional[Tuple[datetime, datetime]]:
        if not isinstance(time_range, str) or '-' not in time_range:
            return None
        parts = [part.strip() for part in time_range.split('-')]
        if len(parts) != 2:
            return None
        try:
            start = datetime.strptime(parts[0], '%H:%M')
            end = datetime.strptime(parts[1], '%H:%M')
            return start, end
        except ValueError:
            return None

    @staticmethod
    def _row_matches_period(row: pd.Series, period_label: str, period_start: datetime, period_end: datetime) -> bool:
        if 'Period' in row.index and pd.notna(row['Period']):
            return str(row['Period']).strip() == period_label
        if 'Time' in row.index and pd.notna(row['Time']):
            parsed = TimetableExporter._parse_time_range(str(row['Time']))
            return parsed == (period_start, period_end)
        return False

    @staticmethod
    def _get_period_match(df: pd.DataFrame, period_label: str, period_start: datetime, period_end: datetime) -> pd.DataFrame:
        if 'Period' in df.columns:
            return df[df['Period'] == period_label]
        if 'Time' in df.columns:
            return df[df.apply(lambda row: TimetableExporter._row_matches_period(row, period_label, period_start, period_end), axis=1)]
        return df.iloc[0:0]

    @staticmethod
    def _get_semester_label(timetable_df: pd.DataFrame) -> str:
        if 'Semester' in timetable_df.columns:
            semesters = timetable_df['Semester'].dropna().unique()
            if len(semesters) == 1:
                return str(semesters[0])
            return ', '.join(sorted(str(s) for s in semesters))
        return 'All'

    @staticmethod
    def _get_branch_label(timetable_df: pd.DataFrame) -> str:
        if 'Branch' in timetable_df.columns:
            branches = timetable_df['Branch'].dropna().unique()
            if len(branches) == 1:
                return str(branches[0])
            return ', '.join(sorted(str(b) for b in branches))
        return 'All'

    @staticmethod
    def _get_export_header_lines(timetable_df: pd.DataFrame) -> List[str]:
        semester_label = TimetableExporter._get_semester_label(timetable_df)
        branch_label = TimetableExporter._get_branch_label(timetable_df)
        return [
            "Vidya Vikas Pratishthan's Institute of Engineering & Technology, Solapur.  NAAC Accredited and ISO 9001 : 2015 Certified Institute",
            "Approved by AICTE , New Delhi, & Govt. of Maharashtra and Affiliated to DBATU, Lonere",
            "DTE CODE: EN6321",
            "72/2 B, Pratapnagar, Soregaon-Dongaon Road, Soregaon, Solapur – 413008    Phone:8380030555    Email:vvpiet@rediffmail.com    Website:ww.vvpengineering.org",
            "Name of Department: ________________________________________________",
            "",
            "Time Table",
            "",
            f"Branch: {branch_label}",
            f"Academic Year: , Semester: {semester_label}, W.E.F.: ",
            ""
        ]

    @staticmethod
    def _get_logo_path() -> Optional[Path]:
        base_dir = Path(__file__).resolve().parent
        possible_names = ['header_1.png', 'logo.png', 'vvp_logo.png', 'logo.jpg', 'vvp_logo.jpg']
        for name in possible_names:
            path = base_dir / name
            if path.exists():
                return path
        return None

    @staticmethod
    def _add_word_header(doc: Document, timetable_df: pd.DataFrame, title: str):
        logo_path = TimetableExporter._get_logo_path()
        if logo_path is not None:
            for section in doc.sections:
                header = section.header
                header.is_linked_to_previous = True
                header_para = header.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = header_para.add_run()
                run.add_picture(str(logo_path), width=Inches(6.5))

        # Add spacing to separate header from title
        doc.add_paragraph()

        # Title centered below header
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title.upper())
        title_run.font.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        semester_label = TimetableExporter._get_semester_label(timetable_df)
        branch_label = TimetableExporter._get_branch_label(timetable_df)
        branch_para = doc.add_paragraph()
        branch_run = branch_para.add_run(f"Branch: {branch_label}")
        branch_run.font.size = Pt(11)
        branch_run.font.bold = True
        branch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_para = doc.add_paragraph()
        info_run = info_para.add_run(f"Academic Year: __________________    Class: ________    Semester: {semester_label}    W.E.F.: __________________")
        info_run.font.size = Pt(11)
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        return doc

    @staticmethod
    def _append_word_note(doc: Document):
        note_para = doc.add_paragraph()
        note_run = note_para.add_run("NOTE:")
        note_run.font.bold = True
        note_run.font.size = Pt(11)
        note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        note1 = doc.add_paragraph()
        note1_run = note1.add_run("10:30 A.M. onwards National Anthem")
        note1_run.font.size = Pt(11)
        note1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        note2 = doc.add_paragraph()
        note2_run = note2.add_run("01.25 P.M. TO 1.30 P.M. Meditation")
        note2_run.font.size = Pt(11)
        note2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return doc

    @staticmethod
    def _get_batch_summary(timetable_df: pd.DataFrame) -> str:
        if 'Batch' not in timetable_df.columns:
            return ''
        batches = sorted({str(b).strip() for b in timetable_df['Batch'].dropna().unique() if str(b).strip() and str(b).strip().lower() != 'all'})
        return ', '.join(batches)

    @staticmethod
    def _append_class_coordinator_section(doc: Document, timetable_df: pd.DataFrame, semester: Optional[str] = None):
        if semester is not None:
            timetable_df = timetable_df[timetable_df['Semester'] == semester]

        batch_summary = TimetableExporter._get_batch_summary(timetable_df)
        if batch_summary:
            line = f"Class Coordinator:-    Proctor Coordinator:-    Batches:- {batch_summary}"
        else:
            line = "Class Coordinator:-    Proctor Coordinator:-    Batches:-"
        coord_para = doc.add_paragraph()
        coord_run = coord_para.add_run(line)
        coord_run.font.size = Pt(11)
        coord_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

        course_rows = timetable_df[['Course Code', 'Course Name', 'Instructor']].drop_duplicates()
        if course_rows.empty:
            return doc

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        headers = ['Sr. No.', 'Course', 'NAME OF THE FACULTY']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, (_, row) in enumerate(course_rows.iterrows(), start=1):
            cells = table.add_row().cells
            cells[0].text = str(idx)
            cells[1].text = f"{row['Course Code']} - {row['Course Name']}"
            cells[2].text = str(row['Instructor'])

        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run('Time Table In-charge               H.O.D.                  Academic Coordinator                   Principal')
        footer_run.font.size = Pt(11)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()
        return doc

    @staticmethod
    def _find_timetable_overlaps(timetable_df: pd.DataFrame) -> List[str]:
        issues = []
        rows = []
        for _, row in timetable_df.iterrows():
            if pd.isna(row.get('Day')) or pd.isna(row.get('Time')):
                continue
            parsed = TimetableExporter._parse_time_range(str(row['Time']))
            if not parsed:
                continue
            start, end = parsed
            rows.append({
                'Day': str(row['Day']).strip(),
                'Semester': str(row.get('Semester', '')).strip(),
                'Instructor': str(row.get('Instructor', '')).strip(),
                'Section': str(row.get('Section', '')).strip(),
                'Batch': str(row.get('Batch', '')).strip(),
                'Course Code': str(row.get('Course Code', '')).strip(),
                'Course Name': str(row.get('Course Name', '')).strip(),
                'Start': start,
                'End': end,
                'Time': str(row['Time']).strip()
            })

        for i in range(len(rows)):
            for j in range(i + 1, len(rows)):
                a = rows[i]
                b = rows[j]
                if a['Day'] != b['Day'] or a['Semester'] != b['Semester']:
                    continue
                if a['Start'] < b['End'] and a['End'] > b['Start']:
                    if a['Instructor'] and a['Instructor'] == b['Instructor']:
                        issues.append(f"Instructor {a['Instructor']} has overlapping classes {a['Course Code']} and {b['Course Code']} on {a['Day']} {a['Time']}")
                    if a['Section'] and a['Section'] == b['Section']:
                        issues.append(f"Section {a['Section']} has overlapping classes {a['Course Code']} and {b['Course Code']} on {a['Day']} {a['Time']}")
                    if a['Batch'] and a['Batch'] != 'All' and a['Batch'] == b['Batch']:
                        issues.append(f"Batch {a['Batch']} has overlapping classes {a['Course Code']} and {b['Course Code']} on {a['Day']} {a['Time']}")

        return sorted(set(issues))

    @staticmethod
    def _append_acknowledgement(doc: Document, issues: List[str]):
        ack_para = doc.add_paragraph()
        if not issues:
            ack_run = ack_para.add_run("ACKNOWLEDGEMENT: There is no overlapping in this timetable.")
            ack_run.font.bold = True
            ack_run.font.color.rgb = RGBColor(0, 128, 0)
            ack_run.font.size = Pt(12)
            ack_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            ack_run = ack_para.add_run("WARNING: Overlaps detected in this timetable. See below details.")
            ack_run.font.bold = True
            ack_run.font.color.rgb = RGBColor(255, 0, 0)
            ack_run.font.size = Pt(12)
            ack_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for issue in issues:
                issue_para = doc.add_paragraph()
                issue_run = issue_para.add_run(issue)
                issue_run.font.size = Pt(11)
                issue_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()
        return doc

    @staticmethod
    def _append_class_coordinator_csv(csv_buffer: BytesIO, timetable_df: pd.DataFrame):
        batch_summary = TimetableExporter._get_batch_summary(timetable_df)
        if batch_summary:
            csv_buffer.write((f'"Class Coordinator:-","Proctor Coordinator:-","Batches:- {batch_summary}"\n').encode('utf-8'))
        else:
            csv_buffer.write((f'"Class Coordinator:-","Proctor Coordinator:-","Batches:-"\n').encode('utf-8'))
        csv_buffer.write(b"\n")
        csv_buffer.write(b'"Sr. No.","Course","NAME OF THE FACULTY"\n')
        course_rows = timetable_df[['Course Code', 'Course Name', 'Instructor']].drop_duplicates()
        for idx, (_, row) in enumerate(course_rows.iterrows(), start=1):
            course_label = f"{row['Course Code']} - {row['Course Name']}"
            csv_buffer.write((f'"{idx}","{course_label}","{row["Instructor"]}"\n').encode('utf-8'))
        csv_buffer.write(b"\n")
        csv_buffer.write(b'"Time Table In-charge","H.O.D.","Academic Coordinator","Principal"\n')
        csv_buffer.write(b"\n")
        return csv_buffer

    @staticmethod
    def _write_header_csv(csv_buffer: BytesIO, header_lines: List[str]):
        for line in header_lines:
            csv_buffer.write((f'"{line.replace("\"", "\"\"")}"\n').encode('utf-8'))

    @staticmethod
    def export_to_csv(timetable_df: pd.DataFrame, config: Optional[dict] = None) -> bytes:
        """Export timetable to CSV format"""
        timetable_df = TimetableExporter._filter_timetable_for_config(timetable_df, config)
        csv_buffer = BytesIO()
        header_lines = TimetableExporter._get_export_header_lines(timetable_df)
        TimetableExporter._write_header_csv(csv_buffer, header_lines)
        timetable_df.to_csv(csv_buffer, index=False)
        csv_buffer.write(b"\n")
        TimetableExporter._append_class_coordinator_csv(csv_buffer, timetable_df)
        TimetableExporter._write_header_csv(csv_buffer, ["NOTE:"])
        TimetableExporter._write_header_csv(csv_buffer, ["10:30 A.M. onwards National Anthem"])
        TimetableExporter._write_header_csv(csv_buffer, ["01.25 P.M. TO 1.30 P.M. Meditation"])
        csv_buffer.seek(0)
        return csv_buffer.getvalue()
    
    @staticmethod
    def export_to_word(timetable_df: pd.DataFrame, title: str = "College Timetable", config: Optional[dict] = None) -> bytes:
        """Export timetable to Word format"""
        timetable_df = TimetableExporter._filter_timetable_for_config(timetable_df, config)
        doc = Document()
        TimetableExporter._add_word_header(doc, timetable_df, "Time Table")

        semesters = sorted([s for s in timetable_df['Semester'].dropna().unique()])
        for sem_index, semester in enumerate(semesters):
            sem_data = timetable_df[timetable_df['Semester'] == semester]
            if sem_data.empty:
                continue
            if sem_index > 0:
                doc.add_page_break()

            sem_heading = doc.add_paragraph()
            sem_run = sem_heading.add_run(f"Semester: {semester}")
            sem_run.font.size = Pt(14)
            sem_run.font.bold = True
            sem_run.font.color.rgb = RGBColor(0, 0, 139)
            sem_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            branch_label = TimetableExporter._get_branch_label(sem_data)
            branch_para = doc.add_paragraph()
            branch_run = branch_para.add_run(f"Branch: {branch_label}")
            branch_run.font.size = Pt(11)
            branch_run.font.bold = True
            branch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            columns = list(sem_data.columns)
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = 'Light Grid Accent 1'

            header_cells = table.rows[0].cells
            for i, column in enumerate(columns):
                header_cells[i].text = str(column)
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            for cell in header_cells:
                cell._element.get_or_add_tcPr().append(shading_elm)

            for _, row in sem_data.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            doc.add_paragraph()
            TimetableExporter._append_class_coordinator_section(doc, sem_data, semester=semester)
            sem_issues = TimetableExporter._find_timetable_overlaps(sem_data)
            TimetableExporter._append_acknowledgement(doc, sem_issues)

        if not semesters:
            no_data_para = doc.add_paragraph()
            no_data_run = no_data_para.add_run("No timetable data available to export.")
            no_data_run.font.size = Pt(12)
            no_data_run.font.bold = True
            no_data_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        return word_buffer.getvalue()

    @staticmethod
    def export_to_word_matrix(timetable_df: pd.DataFrame, title: str = "College Timetable - Timetable Matrix") -> bytes:
        """Export timetable as a matrix format (Time x Days) with course info in cells"""
        doc = Document()
        TimetableExporter._add_word_header(doc, timetable_df, "Time Table - Matrix View")

        semesters = sorted([s for s in timetable_df['Semester'].dropna().unique()])
        for sem_index, semester in enumerate(semesters):
            sem_data = timetable_df[timetable_df['Semester'] == semester]
            if sem_data.empty:
                continue
            if sem_index > 0:
                doc.add_page_break()

            sem_heading = doc.add_paragraph()
            sem_run = sem_heading.add_run(f"Semester: {semester}")
            sem_run.font.size = Pt(14)
            sem_run.font.bold = True
            sem_run.font.color.rgb = RGBColor(0, 0, 139)
            sem_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            branch_label = TimetableExporter._get_branch_label(sem_data)
            branch_para = doc.add_paragraph()
            branch_run = branch_para.add_run(f"Branch: {branch_label}")
            branch_run.font.size = Pt(11)
            branch_run.font.bold = True
            branch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            # Get unique time slots sorted
            time_slots = sorted(sem_data['Time'].dropna().unique(), key=lambda x: (
                datetime.strptime(str(x).split('-')[0].strip(), '%H:%M') if '-' in str(x) else datetime.min
            ))
            
            # Determine which days are present for this semester (include Saturday if used)
            possible_days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
            days = [d for d in possible_days if d in sem_data['Day'].unique()]
            
            # Create matrix: rows = time slots, columns = days
            table = doc.add_table(rows=len(time_slots) + 1, cols=len(days) + 1)
            table.style = 'Light Grid Accent 1'

            # Header row - Day names
            header_cells = table.rows[0].cells
            header_cells[0].text = "Time"
            for paragraph in header_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

            for day_idx, day in enumerate(days):
                header_cells[day_idx + 1].text = day
                for paragraph in header_cells[day_idx + 1].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Shade header row
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            for cell in header_cells:
                cell._element.get_or_add_tcPr().append(shading_elm)

            # Fill in the matrix
            for slot_idx, time_slot in enumerate(time_slots):
                row_cells = table.rows[slot_idx + 1].cells
                row_cells[0].text = str(time_slot)
                for paragraph in row_cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

                # Fill cells for each day
                for day_idx, day in enumerate(days):
                    courses_on_day = sem_data[(sem_data['Day'] == day) & (sem_data['Time'] == time_slot)]
                    cell_text = ""
                    for _, course_row in courses_on_day.iterrows():
                        course_code = str(course_row.get('Course Code', ''))
                        batch = str(course_row.get('Batch', 'All'))
                        if batch and batch != 'All':
                            cell_text += f"{course_code} ({batch})\n"
                        else:
                            cell_text += f"{course_code}\n"
                    
                    row_cells[day_idx + 1].text = cell_text.strip()

            doc.add_paragraph()
            TimetableExporter._append_class_coordinator_section(doc, sem_data, semester=semester)
            sem_issues = TimetableExporter._find_timetable_overlaps(sem_data)
            TimetableExporter._append_acknowledgement(doc, sem_issues)

        if not semesters:
            no_data_para = doc.add_paragraph()
            no_data_run = no_data_para.add_run("No timetable data available to export.")
            no_data_run.font.size = Pt(12)
            no_data_run.font.bold = True
            no_data_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        return word_buffer.getvalue()

    @staticmethod
    def export_day_wise(timetable_df: pd.DataFrame, 
                       title: str = "College Timetable - Day Wise", config: Optional[dict] = None) -> bytes:
        """Export timetable organized by days in Word format"""
        timetable_df = TimetableExporter._filter_timetable_for_config(timetable_df, config)
        doc = Document()
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 139)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Group by day
        days = timetable_df['Day'].unique()
        
        possible_days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
        for day in possible_days:
            if day not in days:
                continue
            
            day_data = timetable_df[timetable_df['Day'] == day]
            if day_data.empty:
                continue
            
            # Add day heading
            day_heading = doc.add_paragraph()
            day_run = day_heading.add_run(f"{day}")
            day_run.font.size = Pt(12)
            day_run.font.bold = True
            day_run.font.color.rgb = RGBColor(0, 0, 139)
            
            # Create table for the day
            table = doc.add_table(rows=1, cols=len(day_data.columns)-1)  # Exclude 'Day' column
            table.style = 'Light Grid Accent 1'
            
            # Add headers (exclude 'Day' column)
            header_cells = table.rows[0].cells
            cols_to_show = [col for col in day_data.columns if col != 'Day']
            for i, column in enumerate(cols_to_show):
                header_cells[i].text = str(column)
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set header background
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            for cell in header_cells:
                cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Add data rows
            for _, row in day_data.iterrows():
                row_cells = table.add_row().cells
                for i, col in enumerate(cols_to_show):
                    row_cells[i].text = str(row[col])
            
            doc.add_paragraph()  # Add space between days
        
        # Save to bytes
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        return word_buffer.getvalue()

    @staticmethod
    def _get_row_period_label(row: pd.Series) -> str:
        if 'Period' in row.index and pd.notna(row['Period']):
            return str(row['Period']).strip()
        if 'Time' in row.index and pd.notna(row['Time']):
            return str(row['Time']).strip()
        return ''

    @staticmethod
    def _filter_timetable_for_config(timetable_df: pd.DataFrame, config: Optional[dict] = None) -> pd.DataFrame:
        """Keep only timetable rows that fall within the configured timings and not inside recess."""
        if timetable_df is None:
            return pd.DataFrame()
        if timetable_df.empty:
            return timetable_df.copy()
        if not config:
            return timetable_df.copy()

        def _coerce_time(value):
            if value is None:
                return datetime.strptime('08:30', '%H:%M')
            if hasattr(value, 'hour') and hasattr(value, 'minute'):
                return datetime.strptime(f"{value.hour:02d}:{value.minute:02d}", '%H:%M')
            return datetime.strptime(str(value), '%H:%M')

        morning_start = _coerce_time(config.get('morning_start', '08:30'))
        morning_end = _coerce_time(config.get('morning_end', '17:00'))
        short_recess_start = _coerce_time(config.get('short_recess_start', '10:30'))
        short_recess_end = _coerce_time(config.get('short_recess_end', '10:45'))
        long_recess_start = _coerce_time(config.get('long_recess_start', '13:00'))
        long_recess_end = _coerce_time(config.get('long_recess_end', '14:00'))

        def _within_config(row):
            parsed = TimetableExporter._parse_time_range(str(row.get('Time', '')).strip())
            if parsed is None:
                return False
            start, end = parsed
            if start < morning_start or end > morning_end:
                return False
            if start < short_recess_end and end > short_recess_start:
                return False
            if start < long_recess_end and end > long_recess_start:
                return False
            return True

        return timetable_df[timetable_df.apply(_within_config, axis=1)].copy()

    @staticmethod
    def export_timetable_matrix_word(timetable_df: pd.DataFrame, title: str = "College Timetable", config: Optional[dict] = None) -> bytes:
        """Export timetable as a matrix-style Word document - Time Slots x Days"""
        doc = Document()
        timetable_df = TimetableExporter._filter_timetable_for_config(timetable_df, config)
        TimetableExporter._add_word_header(doc, timetable_df, "Time Table")

        semesters = sorted(timetable_df['Semester'].dropna().unique())
        for sem_index, semester in enumerate(semesters):
            sem_data = timetable_df[timetable_df['Semester'] == semester].copy()
            if sem_data.empty:
                continue
            if sem_index > 0:
                doc.add_page_break()

            sem_heading = doc.add_paragraph()
            sem_run = sem_heading.add_run(f"Semester: {semester}")
            sem_run.font.size = Pt(14)
            sem_run.font.bold = True
            sem_run.font.color.rgb = RGBColor(0, 0, 139)
            sem_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            branch_label = TimetableExporter._get_branch_label(sem_data)
            branch_para = doc.add_paragraph()
            branch_run = branch_para.add_run(f"Branch: {branch_label}")
            branch_run.font.size = Pt(11)
            branch_run.font.bold = True
            branch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            # Extract unique time slots sorted by start time
            time_slots_raw = sorted(sem_data['Time'].dropna().unique(), key=lambda x: (
                datetime.strptime(str(x).split('-')[0].strip(), '%H:%M') if '-' in str(x) else datetime.min
            ))
            
            # Determine days
            possible_days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
            days = [d for d in possible_days if d in sem_data['Day'].unique()]
            
            # Create table: rows = time slots, columns = days + time slot label
            num_rows = len(time_slots_raw) + 1  # +1 for header
            num_cols = len(days) + 1  # +1 for time slot column
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Light Grid Accent 1'
            table.autofit = False

            # Set column widths
            widths = [Inches(1.2)] + [Inches(1.0)] * len(days)
            for row in table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width

            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Time Slot"
            header_cells[0].paragraphs[0].runs[0].font.bold = True
            header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for day_idx, day in enumerate(days):
                header_cells[day_idx + 1].text = day
                header_cells[day_idx + 1].paragraphs[0].runs[0].font.bold = True
                header_cells[day_idx + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Shade header
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            for cell in header_cells:
                cell._element.get_or_add_tcPr().append(shading_elm)

            # Fill matrix
            for slot_idx, time_slot in enumerate(time_slots_raw):
                row_idx = slot_idx + 1
                row_cells = table.rows[row_idx].cells
                
                # Time slot label
                row_cells[0].text = str(time_slot)
                row_cells[0].paragraphs[0].runs[0].font.bold = True
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # For each day, collect all courses at this time slot
                for day_idx, day in enumerate(days):
                    col_idx = day_idx + 1
                    courses_this_slot = sem_data[(sem_data['Day'] == day) & (sem_data['Time'] == time_slot)]
                    
                    # Clear cell
                    cell = row_cells[col_idx]
                    cell.text = ""
                    
                    # Stack courses in cell
                    if not courses_this_slot.empty:
                        course_labels = []
                        for _, course_row in courses_this_slot.iterrows():
                            label = str(course_row['Course Code'])
                            course_name = str(course_row.get('Course Name', '')).strip()
                            batch = str(course_row.get('Batch', 'All')).strip()
                            if batch and batch != 'All':
                                label += f" ({batch})"
                            if course_name:
                                label += f"\n{course_name}"
                            course_labels.append(label)
                        
                        cell_text = "\n".join(course_labels)
                        cell.text = cell_text
                        
                        # Format text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()
            TimetableExporter._append_class_coordinator_section(doc, sem_data, semester=semester)
            
            # Check overlaps
            sem_issues = TimetableExporter._find_timetable_overlaps(sem_data)
            TimetableExporter._append_acknowledgement(doc, sem_issues)

        if not semesters:
            no_data_para = doc.add_paragraph()
            no_data_run = no_data_para.add_run("No timetable data available to export.")
            no_data_run.font.size = Pt(12)
            no_data_run.font.bold = True
            no_data_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        return word_buffer.getvalue()

    @staticmethod
    def export_timetable_matrix_csv(timetable_df: pd.DataFrame, config: Optional[dict] = None) -> bytes:
        """Export timetable as a matrix-style CSV."""
        timetable_df = TimetableExporter._filter_timetable_for_config(timetable_df, config)
        # Extract unique periods/time slots from timetable data
        if 'Period' in timetable_df.columns and timetable_df['Period'].notna().any():
            # Use Period column if available and populated
            period_labels = [(p, p) for p in sorted(timetable_df['Period'].dropna().unique())]
        else:
            # Fall back to Time column for flexible timing
            unique_times = sorted(timetable_df['Time'].dropna().unique())
            period_labels = [(t, t) for t in unique_times]

        possible_days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
        days = [d for d in possible_days if d in timetable_df['Day'].unique()]

        matrix_rows = []
        for period_label, time_label in period_labels:
            row = {
                'Time Slot': period_label
            }
            for day in days:
                day_rows = timetable_df[timetable_df['Day'] == day]
                
                # Match by Period or Time
                if 'Period' in day_rows.columns and day_rows['Period'].notna().any():
                    period_match = day_rows[day_rows['Period'] == period_label]
                else:
                    period_match = day_rows[day_rows['Time'] == time_label]
                
                if not period_match.empty:
                    values = []
                    for _, item in period_match.iterrows():
                        label = f"{item['Course Code']}"
                        if pd.notna(item['Batch']) and str(item['Batch']) != 'All':
                            label += f" ({item['Batch']})"
                        if pd.notna(item['Branch']) and str(item['Branch']).strip():
                            label += f" {str(item['Branch']).strip()}"
                        label += f" - {item['Course Name']}"
                        values.append(label)
                    row[day] = " | ".join(values)
                else:
                    row[day] = ""
            matrix_rows.append(row)

        matrix_df = pd.DataFrame(matrix_rows)
        csv_buffer = BytesIO()
        header_lines = TimetableExporter._get_export_header_lines(timetable_df)
        TimetableExporter._write_header_csv(csv_buffer, header_lines)
        matrix_df.to_csv(csv_buffer, index=False)
        csv_buffer.write(b"\n")
        TimetableExporter._write_header_csv(csv_buffer, ["NOTE:"])
        TimetableExporter._write_header_csv(csv_buffer, ["10:30 A.M. onwards National Anthem"])
        TimetableExporter._write_header_csv(csv_buffer, ["01.25 P.M. TO 1.30 P.M. Meditation"])
        csv_buffer.seek(0)
        return csv_buffer.getvalue()
